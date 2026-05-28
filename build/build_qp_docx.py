from __future__ import annotations

import argparse
import hashlib
import json
import re
from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from docx import Document


BASE = Path(__file__).resolve().parent
PROJECT_ROOT = BASE.parent
TEMPLATE = PROJECT_ROOT / "template" / "樣本.docx"
SOURCE_MD = PROJECT_ROOT / "doc" / "L2-01_vulnerability-handling-and-disclosure-process.md"
FLOW_IMAGE = PROJECT_ROOT / "template" / "vul_handle_n_disclose_flow.png"
OUTPUT = PROJECT_ROOT / "doc" / "QP-30-01 事件處理程序 V1.0.docx"
TRANSLATION_DIR = PROJECT_ROOT / "prompt" / "translation"
GLOSSARY = TRANSLATION_DIR / "glossary_psirt.json"
BLACKLIST = TRANSLATION_DIR / "blacklist_english.json"
TRANSLATION_MEMORY_FILE = TRANSLATION_DIR / "translation_memory_l2_01.json"
FENCE = chr(96) * 3
ARGOS_READY: bool | None = None
SOURCE_START_HEADING = "## 1. 目的 Purpose"
SOURCE_START_RE = re.compile(r"^##\s+(?:\d+\.\s+)?目的\s+Purpose\s*$")
SOURCE_CONTROL_HEADING = "# L2-01：弱點處理與揭露程序 Vulnerability Handling and Disclosure Process"
KNOWN_MERMAID_SHA256 = "cf06288f248d0a26ba8b71655d84ed9d9e6f947ba8d7787b1e71bf52a901cf27"
MISSING_TRANSLATION_REPORT = PROJECT_ROOT / "tmp" / "missing_translations.txt"
BLOCKED_ENGLISH_REPORT = PROJECT_ROOT / "tmp" / "blocked_english.txt"
TRANSLATION_MEMORY: dict[str, str] = {}
BLOCKED_PHRASES: list[dict[str, str]] = []
MISSING_TRANSLATIONS: set[str] = set()
BLOCKED_ENGLISH: list[str] = []


def resolve_output_path(path: Path, date_suffix: str | None = None) -> Path:
    if not path.exists():
        return path

    suffix = date_suffix or datetime.now().strftime("%Y%m%d")
    dated = path.with_name(f"{path.stem}_{suffix}{path.suffix}")
    if not dated.exists():
        return dated

    version = 2
    while True:
        candidate = path.with_name(f"{path.stem}_{suffix}_v{version}{path.suffix}")
        if not candidate.exists():
            return candidate
        version += 1


def validate_inputs() -> list[str]:
    missing = []
    for label, path in (
        ("Word 樣本", TEMPLATE),
        ("來源 Markdown", SOURCE_MD),
        ("流程圖圖片", FLOW_IMAGE),
        ("術語表", GLOSSARY),
        ("英文黑名單", BLACKLIST),
        ("人工翻譯記憶庫", TRANSLATION_MEMORY_FILE),
    ):
        if not path.exists():
            missing.append(f"{label} 不存在: {path}")
    return missing


TERM_MAP = {
    "弱點": "vulnerability",
    "漏洞": "vulnerability",
    "處理": "handling",
    "揭露": "disclosure",
    "程序": "process",
    "目的": "Purpose",
    "範圍": "Scope",
    "適用性": "Applicability",
    "風險": "risk",
    "加嚴要求": "enhanced requirements",
    "程序原則": "Governing Principles",
    "準備能力": "Preparedness Capabilities",
    "名詞": "Definitions",
    "角色": "Roles",
    "利害關係人": "stakeholders",
    "程序總覽": "Process Overview",
    "流程要求": "Process Requirements",
    "受理": "intake",
    "建案": "registration",
    "初步回覆": "initial response",
    "受理判定": "acceptance decision",
    "驗證": "validation",
    "產品影響分析": "product impact assessment",
    "已遭利用": "exploited",
    "狀態判定": "status determination",
    "修補": "remediation",
    "緩解": "mitigation",
    "規劃": "planning",
    "時程基準": "timing baseline",
    "發布": "release",
    "交付": "delivery",
    "通知": "notification",
    "決策": "decision",
    "強制通報": "mandatory reporting",
    "結案": "closure",
    "持續改善": "continual improvement",
    "對外溝通政策": "External Communication Policy",
    "對外窗口": "external contact point",
    "通道": "channel",
    "時限政策": "timeline policy",
    "對外訊息": "external message",
    "最小內容": "minimum content",
    "禁止過早揭露內容": "prohibited premature disclosure content",
    "利害關係人管理": "Stakeholder Management",
    "記錄": "recordkeeping",
    "資料保護": "data protection",
    "必填紀錄": "required records",
    "附件": "attachments",
    "敏感資料": "sensitive data",
    "保留": "retention",
    "稽核": "audit",
    "量測": "metrics",
    "文件治理": "document governance",
    "版控": "version control",
    "修訂紀錄": "revision history",
}


EXACT_TRANSLATIONS = {
    "文件屬性：程序（Level 2）": "Document attribute: Procedure (Level 2)",
    "適用標準：ISO/IEC 29147:2018、ISO/IEC 30111:2019、prEN 40000-1-3:2026 draft、CRA（Regulation (EU) 2024/2847）相關要求": "Applicable standards: ISO/IEC 29147:2018, ISO/IEC 30111:2019, prEN 40000-1-3:2026 draft, and related CRA (Regulation (EU) 2024/2847) requirements",
    "版本：v0.41（草案）": "Version: v0.41 (Draft)",
    "修訂者：Stanley Huang": "Revised by: Stanley Huang",
    "文件擁有者：IEI PSIRT": "Document owner: IEI PSIRT",
    "建立 IEI 產品弱點受理、驗證、分析、修補、揭露、通報與結案之統一流程，確保：": "Establish a unified process for IEI product vulnerability intake, verification, analysis, remediation, disclosure, reporting, and closure to ensure that:",
    "本程序適用於下列與 IEI 產品或其數位元件相關之弱點案件：": "This procedure applies to the following vulnerability cases related to IEI products or their digital elements:",
    "本程序不取代一般 IT 事故處理程序；若案件涉及企業內部 IT 環境而非產品本身，應轉依適用之資訊安全事件程序辦理。": "This procedure does not replace the general IT incident handling procedure; if a case involves the corporate IT environment rather than the product itself, it shall be handled under the applicable information security incident procedure.",
    "對外之弱點受理與協調揭露由 PSIRT 統一窗口處理，避免重複承諾或口徑不一致。": "External vulnerability intake and coordinated disclosure shall be handled by PSIRT through a single point of contact to avoid duplicate commitments or inconsistent messaging.",
    "在修補或風險控制完成前，僅共享完成驗證、修補、通知與法遵所需之最小資訊，避免過早揭露可武器化細節。": "Before remediation or risk control is completed, only the minimum information necessary for verification, remediation, notification, and compliance shall be shared to avoid premature disclosure of weaponizable details.",
    "所有關鍵判定、核准、例外、對外溝通摘要與證據連結均應記錄於 Jira 主案件中。": "All key determinations, approvals, exceptions, external communication summaries, and evidence links shall be recorded in the master Jira case.",
    "若案件可對產品造成重大安全衝擊或疑似已遭利用，PSIRT 應立即升級為高優先案件，不受一般分析排程限制。": "If a case may cause significant security impact to a product or is suspected to have been exploited, PSIRT shall immediately escalate it as a high-priority case outside the normal analysis schedule.",
}


SECTION_TRANSLATIONS = {
    "L2-01：弱點處理與揭露程序 Vulnerability Handling and Disclosure Process": "L2-01: Vulnerability Handling and Disclosure Process",
    "目的 Purpose": "Purpose",
    "1. 目的 Purpose": "1. Purpose",
    "範圍 Scope": "Scope",
    "2. 範圍 Scope": "2. Scope",
    "適用性與風險式加嚴要求 Applicability and Risk-based Enhancements": "Applicability and Risk-based Enhancements",
    "2.1 適用性與風險式加嚴要求 Applicability and Risk-based Enhancements": "2.1 Applicability and Risk-based Enhancements",
    "程序原則 Governing Principles": "Governing Principles",
    "3. 程序原則 Governing Principles": "3. Governing Principles",
    "準備能力 Preparedness Capabilities": "Preparedness Capabilities",
    "3.1 準備能力 Preparedness Capabilities": "3.1 Preparedness Capabilities",
    "名詞 Definitions": "Definitions",
    "4. 名詞 Definitions": "4. Definitions",
    "弱點案件 Vulnerability Case": "Vulnerability Case",
    "4.1 弱點案件 Vulnerability Case": "4.1 Vulnerability Case",
    "協調式弱點揭露 CVD": "Coordinated Vulnerability Disclosure (CVD)",
    "4.2 協調式弱點揭露 CVD": "4.2 Coordinated Vulnerability Disclosure (CVD)",
    "已遭主動利用弱點 Actively Exploited Vulnerability": "Actively Exploited Vulnerability",
    "4.3 已遭主動利用弱點 Actively Exploited Vulnerability": "4.3 Actively Exploited Vulnerability",
    "嚴重事件 Severe Incident": "Severe Incident",
    "4.4 嚴重事件 Severe Incident": "4.4 Severe Incident",
    "支援期間 Support Period": "Support Period",
    "4.5 支援期間 Support Period": "4.5 Support Period",
    "Jira 主案件 Master Jira Record": "Master Jira Record",
    "4.6 Jira 主案件 Master Jira Record": "4.6 Master Jira Record",
    "角色與利害關係人 Roles and Stakeholders": "Roles and Stakeholders",
    "5. 角色與利害關係人 Roles and Stakeholders": "5. Roles and Stakeholders",
    "程序總覽 Process Overview": "Process Overview",
    "6. 程序總覽 Process Overview": "6. Process Overview",
    "流程要求 Process Requirements": "Process Requirements",
    "7. 流程要求 Process Requirements": "7. Process Requirements",
    "受理與建案 Intake and Registration": "Intake and Registration",
    "7.1 受理與建案 Intake and Registration": "7.1 Intake and Registration",
    "初步回覆與受理判定 Initial Response and Acceptance": "Initial Response and Acceptance",
    "7.2 初步回覆與受理判定 Initial Response and Acceptance": "7.2 Initial Response and Acceptance",
    "驗證與產品影響分析 Verification and Product Impact Assessment": "Verification and Product Impact Assessment",
    "7.3 驗證與產品影響分析 Verification and Product Impact Assessment": "7.3 Verification and Product Impact Assessment",
    "已遭利用狀態判定 Exploitation Status Determination": "Exploitation Status Determination",
    "7.4 已遭利用狀態判定 Exploitation Status Determination": "7.4 Exploitation Status Determination",
    "修補或緩解規劃 Remediation and Mitigation Planning": "Remediation and Mitigation Planning",
    "7.5 修補或緩解規劃 Remediation and Mitigation Planning": "7.5 Remediation and Mitigation Planning",
    "修補時程基準 Remediation Timing Baseline": "Remediation Timing Baseline",
    "7.6 修補時程基準 Remediation Timing Baseline": "7.6 Remediation Timing Baseline",
    "驗證、發布與交付 Validation, Release and Delivery": "Validation, Release and Delivery",
    "7.7 驗證、發布與交付 Validation, Release and Delivery": "7.7 Validation, Release and Delivery",
    "揭露與通知決策 Disclosure and Notification Decision": "Disclosure and Notification Decision",
    "7.8 揭露與通知決策 Disclosure and Notification Decision": "7.8 Disclosure and Notification Decision",
    "CRA 強制通報 CRA Mandatory Reporting": "CRA Mandatory Reporting",
    "7.9 CRA 強制通報 CRA Mandatory Reporting": "7.9 CRA Mandatory Reporting",
    "結案與持續改善 Closure and Continuous Improvement": "Closure and Continual Improvement",
    "7.10 結案與持續改善 Closure and Continuous Improvement": "7.10 Closure and Continual Improvement",
    "對外溝通政策 External Communication Policy": "External Communication Policy",
    "8. 對外溝通政策 External Communication Policy": "8. External Communication Policy",
    "對外窗口與通道": "External Contact Points and Channels",
    "8.1 對外窗口與通道": "8.1 External Contact Points and Channels",
    "時限政策": "Timeline Policy",
    "8.2 時限政策": "8.2 Timeline Policy",
    "對外訊息最小內容": "Minimum External Message Content",
    "8.3 對外訊息最小內容": "8.3 Minimum External Message Content",
    "禁止過早揭露內容": "Prohibited Premature Disclosure Content",
    "8.4 禁止過早揭露內容": "8.4 Prohibited Premature Disclosure Content",
    "利害關係人管理 Stakeholder Management": "Stakeholder Management",
    "9. 利害關係人管理 Stakeholder Management": "9. Stakeholder Management",
    "Jira 記錄與資料保護 Jira Recordkeeping and Data Protection": "Jira Recordkeeping and Data Protection",
    "10. Jira 記錄與資料保護 Jira Recordkeeping and Data Protection": "10. Jira Recordkeeping and Data Protection",
    "必填紀錄": "Required Records",
    "10.1 必填紀錄": "10.1 Required Records",
    "已遭利用與 CRA 判定留痕要求": "Traceability Requirements for Exploitation and CRA Determinations",
    "10.2 已遭利用與 CRA 判定留痕要求": "10.2 Traceability Requirements for Exploitation and CRA Determinations",
    "附件與敏感資料": "Attachments and Sensitive Data",
    "10.3 附件與敏感資料": "10.3 Attachments and Sensitive Data",
    "保留與稽核": "Retention and Audit",
    "10.4 保留與稽核": "10.4 Retention and Audit",
    "Jira 資料保存政策": "Jira Data Retention Policy",
    "10.5 Jira 資料保存政策": "10.5 Jira Data Retention Policy",
    "稽核與量測 Audit and Metrics": "Audit and Metrics",
    "11. 稽核與量測 Audit and Metrics": "11. Audit and Metrics",
    "文件治理與版控 Document Control": "Document Control",
    "12. 文件治理與版控 Document Control": "12. Document Control",
    "修訂紀錄 Revision History": "Revision History",
    "13. 修訂紀錄 Revision History": "13. Revision History",
}


def has_cjk(text: str) -> bool:
    return any("\u4e00" <= ch <= "\u9fff" for ch in text)


def clean_inline(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    return text.strip().strip(chr(96))


def load_json_object(path: Path) -> dict:
    with path.open(encoding="utf-8") as handle:
        data = json.load(handle)
    if not isinstance(data, dict):
        raise SystemExit(f"JSON 格式錯誤，根節點必須為 object: {path}")
    return data


def load_translation_memory() -> dict[str, str]:
    data = load_json_object(TRANSLATION_MEMORY_FILE)
    translations = data.get("translations", {})
    if not isinstance(translations, dict):
        raise SystemExit(f"translation memory 格式錯誤: {TRANSLATION_MEMORY_FILE}")

    memory: dict[str, str] = {}
    for zh, en in translations.items():
        source = clean_inline(str(zh))
        target = clean_inline(str(en))
        if not source or not target:
            continue
        if has_cjk(target):
            raise SystemExit(f"translation memory 英文含中文，請修正: {source}")
        memory[source] = target
    return memory


def validate_translation_memory(memory: dict[str, str]) -> None:
    for source, target in memory.items():
        validate_english_text(target, source)


def load_blocked_phrases() -> list[dict[str, str]]:
    blocked: list[dict[str, str]] = []

    blacklist = load_json_object(BLACKLIST)
    for entry in blacklist.get("blocked_phrases", []):
        if isinstance(entry, dict) and entry.get("phrase"):
            blocked.append(
                {
                    "phrase": str(entry["phrase"]),
                    "reason": str(entry.get("reason", "")),
                }
            )

    glossary = load_json_object(GLOSSARY)
    for entry in glossary.get("forbidden_terms", []):
        if isinstance(entry, dict) and entry.get("term"):
            blocked.append(
                {
                    "phrase": str(entry["term"]),
                    "reason": str(entry.get("reason", "Forbidden by glossary.")),
                }
            )
    return blocked


def record_blocked_english(text: str, source: str, phrase: str, reason: str) -> None:
    BLOCKED_ENGLISH.append(f"{phrase} | {reason} | source={source} | text={text}")


def validate_english_text(text: str, source: str) -> bool:
    lowered = text.lower()
    ok = True
    for entry in BLOCKED_PHRASES:
        phrase = entry["phrase"]
        if phrase.lower() in lowered:
            record_blocked_english(text, source, phrase, entry.get("reason", ""))
            ok = False
    return ok


def approved_translation(source: str, target: str) -> str:
    if validate_english_text(target, source):
        return target
    return ""


def mermaid_blocks(lines: list[str]) -> list[str]:
    blocks = []
    in_code = False
    code_lang = ""
    code_buf: list[str] = []
    for raw in lines:
        line = raw.strip()
        if line.startswith(FENCE):
            if not in_code:
                in_code = True
                code_lang = line[3:].strip()
                code_buf = []
            else:
                if code_lang == "mermaid":
                    blocks.append("\n".join(code_buf).strip() + "\n")
                in_code = False
                code_lang = ""
                code_buf = []
            continue
        if in_code:
            code_buf.append(raw)
    return blocks


def validate_source_lines(lines: list[str]) -> None:
    if not any(SOURCE_START_RE.match(line.strip()) for line in lines):
        raise SystemExit(
            "找不到正式內容起始標題。預期格式為 '## 1. 目的 Purpose' 或 '## 目的 Purpose'。"
        )

    blocks = mermaid_blocks(lines)
    if not blocks:
        return
    if len(blocks) != 1:
        raise SystemExit(f"Mermaid 區塊數量為 {len(blocks)}，需人工確認流程圖圖片對應關係。")

    digest = hashlib.sha256(blocks[0].encode("utf-8")).hexdigest()
    if digest != KNOWN_MERMAID_SHA256:
        raise SystemExit(
            "來源 Mermaid 流程圖內容已變更，請先更新 template/vul_handle_n_disclose_flow.png "
            "並同步 KNOWN_MERMAID_SHA256。"
        )


def trim_source_lines(lines: list[str]) -> list[str]:
    for idx, line in enumerate(lines):
        if SOURCE_START_RE.match(line.strip()):
            return lines[idx:]
    raise SystemExit(
        "找不到正式內容起始標題。預期格式為 '## 1. 目的 Purpose' 或 '## 目的 Purpose'。"
    )


def strip_heading_number(text: str) -> str:
    return re.sub(r"^(?:[1-9]\d?)(?:\.(?:[1-9]\d?))*\.?\s+", "", clean_inline(text)).strip()


def write_missing_translation_report() -> None:
    MISSING_TRANSLATION_REPORT.parent.mkdir(parents=True, exist_ok=True)
    content = ["# Missing translations", ""]
    content.extend(f"- {item}" for item in sorted(MISSING_TRANSLATIONS))
    MISSING_TRANSLATION_REPORT.write_text("\n".join(content) + "\n", encoding="utf-8")


def write_blocked_english_report() -> None:
    BLOCKED_ENGLISH_REPORT.parent.mkdir(parents=True, exist_ok=True)
    content = ["# Blocked English", ""]
    content.extend(f"- {item}" for item in BLOCKED_ENGLISH)
    BLOCKED_ENGLISH_REPORT.write_text("\n".join(content) + "\n", encoding="utf-8")


def split_markdown_row(raw: str) -> list[str]:
    cells = []
    current = []
    escaped = False
    for ch in raw.strip().strip("|"):
        if escaped:
            current.append(ch)
            escaped = False
            continue
        if ch == "\\":
            escaped = True
            continue
        if ch == "|":
            cells.append(clean_inline("".join(current).strip()))
            current = []
            continue
        current.append(ch)
    cells.append(clean_inline("".join(current).strip()))
    return cells


def translate_sentence(text: str) -> str:
    source = clean_inline(text)
    if source in TRANSLATION_MEMORY:
        return approved_translation(source, TRANSLATION_MEMORY[source])
    if source in EXACT_TRANSLATIONS:
        return approved_translation(source, EXACT_TRANSLATIONS[source])
    if has_cjk(source):
        MISSING_TRANSLATIONS.add(source)
        return ""
    return approved_translation(source, source)


def translate(text: str) -> str:
    text = clean_inline(text)
    if not text:
        return ""
    if text in TRANSLATION_MEMORY:
        return approved_translation(text, TRANSLATION_MEMORY[text])
    if text in SECTION_TRANSLATIONS:
        return approved_translation(text, SECTION_TRANSLATIONS[text])
    if text in EXACT_TRANSLATIONS:
        return approved_translation(text, EXACT_TRANSLATIONS[text])
    if not has_cjk(text):
        return approved_translation(text, text)
    if "：" in text:
        head, tail = text.split("：", 1)
        head_en = translate(head)
        tail_en = translate_sentence(tail)
        if not head_en or not tail_en:
            MISSING_TRANSLATIONS.add(text)
            return ""
        return approved_translation(text, f"{head_en}: {tail_en}")
    return translate_sentence(text)


def remove_body_content(doc: "Document") -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


PARAGRAPH_FORMAT_ATTRS = (
    "alignment",
    "first_line_indent",
    "keep_together",
    "keep_with_next",
    "left_indent",
    "line_spacing",
    "line_spacing_rule",
    "page_break_before",
    "right_indent",
    "space_after",
    "space_before",
    "widow_control",
)


def resolve_paragraph_format_value(paragraph, attr: str):
    value = getattr(paragraph.paragraph_format, attr)
    if value is not None:
        return value
    style = paragraph.style
    while style is not None:
        value = getattr(style.paragraph_format, attr)
        if value is not None:
            return value
        style = style.base_style
    return None


def sync_paragraph_format(source, target) -> None:
    for attr in PARAGRAPH_FORMAT_ATTRS:
        value = resolve_paragraph_format_value(source, attr)
        if value is not None:
            setattr(target.paragraph_format, attr, value)


def add_para(doc: "Document", text: str, style: str = "Body"):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def english_style(zh_style: str) -> str:
    if zh_style in {"H1", "H2", "H3", "H4", "H5", "Body"}:
        return f"{zh_style}EN"
    return "BodyEN"


def add_bilingual(doc: "Document", zh: str, zh_style: str, en_style: str | None = None) -> None:
    zh = clean_inline(zh)
    if not zh:
        return
    en = translate(zh)
    zh_para = add_para(doc, zh, zh_style)
    if en and en != zh:
        en_para = add_para(doc, en, en_style or english_style(zh_style))
        sync_paragraph_format(zh_para, en_para)


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = split_markdown_row(lines[i])
        if not all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_table(doc: "Document", rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=width)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for cidx in range(width):
            text = row[cidx] if cidx < len(row) else ""
            cell = cells[cidx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            if ridx == 0:
                run.bold = True
            en = translate(text)
            if en and en != text:
                en_para = cell.add_paragraph(en)
                sync_paragraph_format(p, en_para)


def render_markdown(doc: "Document", lines: list[str]) -> None:
    from docx.enum.text import WD_BREAK
    from docx.shared import Inches

    in_code = False
    code_lang = ""
    code_buf: list[str] = []
    i = 0
    while i < len(lines):
        raw = lines[i].rstrip("\n")
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line.startswith(FENCE):
            if not in_code:
                in_code = True
                code_lang = line[3:].strip()
                code_buf = []
            else:
                if code_lang == "mermaid" and FLOW_IMAGE.exists():
                    add_para(doc, "流程圖", "H2")
                    add_para(doc, "Flow Chart", "H2EN")
                    doc.add_picture(str(FLOW_IMAGE), width=Inches(6.3))
                else:
                    add_para(doc, "\n".join(code_buf), "HTML")
                in_code = False
                code_lang = ""
                code_buf = []
            i += 1
            continue
        if in_code:
            code_buf.append(raw)
            i += 1
            continue
        if line.startswith("|"):
            rows, next_i = parse_markdown_table(lines, i)
            add_table(doc, rows)
            i = next_i
            continue
        m = re.match(r"^(#{1,5})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            style = "H1" if level <= 2 else "H2" if level == 3 else "H3"
            add_bilingual(doc, strip_heading_number(m.group(2)), style)
            i += 1
            continue
        m = re.match(r"^[-*]\s+(.+)$", line)
        if m:
            add_bilingual(doc, m.group(1), "af", "BodyEN")
            i += 1
            continue
        m = re.match(r"^\d+\.\s+(.+)$", line)
        if m:
            add_bilingual(doc, m.group(1), "af", "BodyEN")
            i += 1
            continue
        if line == "---":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            i += 1
            continue
        add_bilingual(doc, line, "Body")
        i += 1


def set_default_fonts(doc: "Document") -> None:
    from docx.shared import Pt

    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.size is None:
                run.font.size = Pt(10.5)


def main() -> None:
    global TRANSLATION_MEMORY
    global BLOCKED_PHRASES

    parser = argparse.ArgumentParser(description="Build QP DOCX from the L2 Markdown source.")
    parser.add_argument("--dry-run", action="store_true", help="Show resolved paths without writing the DOCX.")
    parser.add_argument("--date-suffix", help="Override the YYYYMMDD suffix used when the output exists.")
    args = parser.parse_args()

    output = resolve_output_path(OUTPUT, args.date_suffix)
    missing = validate_inputs()
    if missing:
        for message in missing:
            print(message)
        raise SystemExit(1)

    raw_lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    validate_source_lines(raw_lines)
    BLOCKED_PHRASES = load_blocked_phrases()
    TRANSLATION_MEMORY = load_translation_memory()
    validate_translation_memory(TRANSLATION_MEMORY)

    if BLOCKED_ENGLISH:
        write_blocked_english_report()
        raise SystemExit(
            f"blocked_english={len(BLOCKED_ENGLISH)}; "
            f"report={BLOCKED_ENGLISH_REPORT}"
        )

    if args.dry_run:
        print(f"source={SOURCE_MD}")
        print(f"template={TEMPLATE}")
        print(f"flow_image={FLOW_IMAGE}")
        print(f"requested_output={OUTPUT}")
        print(f"resolved_output={output}")
        print(f"output_exists={OUTPUT.exists()}")
        print("write=false")
        return

    from docx import Document

    doc = Document(str(TEMPLATE))
    remove_body_content(doc)
    lines = trim_source_lines(raw_lines)
    render_markdown(doc, lines)
    if BLOCKED_ENGLISH:
        write_blocked_english_report()
    if MISSING_TRANSLATIONS:
        write_missing_translation_report()
    if MISSING_TRANSLATIONS or BLOCKED_ENGLISH:
        details = []
        if MISSING_TRANSLATIONS:
            details.append(f"missing_translations={len(MISSING_TRANSLATIONS)}; report={MISSING_TRANSLATION_REPORT}")
        if BLOCKED_ENGLISH:
            details.append(f"blocked_english={len(BLOCKED_ENGLISH)}; report={BLOCKED_ENGLISH_REPORT}")
        raise SystemExit(
            "; ".join(details)
        )
    set_default_fonts(doc)
    doc.save(str(output))
    print(output)


if __name__ == "__main__":
    main()
